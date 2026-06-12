#!/usr/bin/env python3
"""
Build DOCX fixtures for the spoor test suite.

Each fixture targets a specific category of behavior the extractor must handle.
We use python-docx for natural docs and zipfile+raw XML for edge cases that
python-docx can't easily produce (custom namespaces, malformed structures).
"""
from pathlib import Path
import zipfile

from docx import Document
from docx.shared import Inches

OUT = Path(__file__).resolve().parent.parent / "fixtures" / "docx"
OUT.mkdir(parents=True, exist_ok=True)


# ---------- 01: basic structure (headings, paragraphs, inline formatting) ----------
def build_01_basic():
    d = Document()
    d.add_heading("First Heading", level=1)
    d.add_paragraph("A normal paragraph.")
    d.add_heading("Second Heading", level=2)
    p = d.add_paragraph()
    p.add_run("plain ")
    p.add_run("bold ").bold = True
    p.add_run("italic ").italic = True
    r = p.add_run("bold-italic")
    r.bold = True
    r.italic = True
    d.add_heading("Third level", level=3)
    d.add_paragraph("Trailing text.")
    d.save(OUT / "01_basic.docx")


# ---------- 02: lists (bullets, numbers, nested) ----------
def build_02_lists():
    d = Document()
    d.add_heading("Lists", 1)
    d.add_paragraph("Bullet 1", style="List Bullet")
    d.add_paragraph("Bullet 2", style="List Bullet")
    d.add_paragraph("Number 1", style="List Number")
    d.add_paragraph("Number 2", style="List Number")
    d.add_paragraph("Number 3", style="List Number")
    d.save(OUT / "02_lists.docx")


# ---------- 03: tables ----------
def build_03_tables():
    d = Document()
    d.add_heading("Tables", 1)
    t = d.add_table(rows=3, cols=3)
    headers = t.rows[0].cells
    headers[0].text = "Name"
    headers[1].text = "Score"
    headers[2].text = "Note"
    rows = [("Alice", "95", "first"), ("Bob", "80", "with | pipe")]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    d.save(OUT / "03_tables.docx")


# ---------- 04: hyperlinks (raw XML, python-docx hyperlinks are awkward) ----------
def build_04_hyperlinks():
    document_xml = """<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<w:body>
<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t>Hyperlinks</w:t></w:r></w:p>
<w:p><w:r><w:t xml:space="preserve">See </w:t></w:r>
<w:hyperlink r:id="rIdL1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>example</w:t></w:r></w:hyperlink>
<w:r><w:t xml:space="preserve"> and </w:t></w:r>
<w:hyperlink r:id="rIdL2"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>another</w:t></w:r></w:hyperlink>
<w:r><w:t>.</w:t></w:r></w:p>
</w:body></w:document>"""
    write_minimal_docx(
        OUT / "04_hyperlinks.docx",
        document_xml,
        extra_rels=[
            ('rIdL1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'https://example.com'),
            ('rIdL2', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 'https://another.org/path?q=1'),
        ],
    )


# ---------- 05: footnotes ----------
def build_05_footnotes():
    document_xml = """<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p><w:r><w:t xml:space="preserve">A claim with citation</w:t></w:r>
<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="1"/></w:r>
<w:r><w:t xml:space="preserve"> and another</w:t></w:r>
<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="2"/></w:r>
<w:r><w:t>.</w:t></w:r></w:p>
</w:body></w:document>"""
    footnotes_xml = """<?xml version="1.0"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:footnote w:id="1"><w:p><w:r><w:t>First footnote source.</w:t></w:r></w:p></w:footnote>
<w:footnote w:id="2"><w:p><w:r><w:t>Second source, see appendix.</w:t></w:r></w:p></w:footnote>
</w:footnotes>"""
    write_minimal_docx(
        OUT / "05_footnotes.docx",
        document_xml,
        footnotes_xml=footnotes_xml,
    )


# ---------- 06: unicode, smart quotes, RTL, emoji ----------
def build_06_unicode():
    d = Document()
    d.add_heading("中文标题 / 日本語", 1)
    d.add_paragraph("中英 mixed text with 한글 そして emoji 🎉🚀")
    d.add_paragraph("RTL: עברית and العربية")
    d.add_paragraph("Math: α β γ ∑ ∫ ≈ ≠ ≤")
    d.add_paragraph("Smart: \u201chello\u201d \u2018world\u2019 — em-dash")
    d.save(OUT / "06_unicode.docx")


# ---------- 07: custom XML namespace prefix (robustness test) ----------
def build_07_custom_prefix():
    """extract-text passes this; ours must too — match by namespace URI, not prefix."""
    document_xml = """<?xml version="1.0"?>
<x:document xmlns:x="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<x:body>
<x:p><x:pPr><x:pStyle x:val="Heading1"/></x:pPr><x:r><x:t>Custom prefix x</x:t></x:r></x:p>
<x:p><x:r><x:t>Plain paragraph under custom prefix.</x:t></x:r></x:p>
</x:body></x:document>"""
    write_minimal_docx(OUT / "07_custom_prefix.docx", document_xml)


# ---------- 08: empty document ----------
def build_08_empty():
    d = Document()
    d.save(OUT / "08_empty.docx")


# ---------- 09: only whitespace and empty paragraphs ----------
def build_09_whitespace():
    d = Document()
    d.add_paragraph("")
    d.add_paragraph("   ")
    d.add_paragraph("\t")
    d.add_paragraph("real content")
    d.add_paragraph("")
    d.save(OUT / "09_whitespace.docx")


# ---------- 10: heading-level cap (Heading 7+ should fall back) ----------
def build_10_heading_levels():
    d = Document()
    d.add_heading("H1", 1)
    d.add_heading("H2", 2)
    d.add_heading("H3", 3)
    d.add_heading("H4", 4)
    d.add_heading("H5", 5)
    d.add_heading("H6", 6)
    d.save(OUT / "10_heading_levels.docx")


# ---------- 11: text run with leading/trailing whitespace (xml:space) ----------
def build_11_whitespace_runs():
    document_xml = """<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p><w:r><w:t xml:space="preserve">Leading </w:t></w:r>
<w:r><w:rPr><w:b/></w:rPr><w:t>middle</w:t></w:r>
<w:r><w:t xml:space="preserve"> trailing.</w:t></w:r></w:p>
<w:p><w:r><w:t>NoSpaces</w:t></w:r><w:r><w:t>Concatenated</w:t></w:r></w:p>
</w:body></w:document>"""
    write_minimal_docx(OUT / "11_whitespace_runs.docx", document_xml)


# ---------- 12: tracked changes (we should accept the insertions and skip deletions) ----------
def build_12_tracked_changes():
    document_xml = """<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p>
<w:r><w:t xml:space="preserve">The price is </w:t></w:r>
<w:del w:id="1" w:author="A" w:date="2025-01-01T00:00:00Z"><w:r><w:delText>30</w:delText></w:r></w:del>
<w:ins w:id="2" w:author="A" w:date="2025-01-01T00:00:00Z"><w:r><w:t>60</w:t></w:r></w:ins>
<w:r><w:t xml:space="preserve"> dollars.</w:t></w:r>
</w:p>
</w:body></w:document>"""
    write_minimal_docx(OUT / "12_tracked_changes.docx", document_xml)


# ---------- 13: bold/italic/hyperlink on whitespace-only runs (must not panic) ----------
def build_13_formatted_whitespace_only_runs():
    """Real Word docs often put w:b / w:i / hyperlink on runs that are only spaces or w:br."""
    document_xml = """<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<w:body>
<w:p><w:r><w:t>intro</w:t></w:r>
<w:r><w:rPr><w:b/></w:rPr><w:t xml:space="preserve"> </w:t></w:r>
<w:r><w:t>after space</w:t></w:r></w:p>
<w:p><w:r><w:rPr><w:i/></w:rPr><w:br/></w:r><w:r><w:t>line2</w:t></w:r></w:p>
<w:p>
<w:hyperlink r:id="rIdL1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t xml:space="preserve"> </w:t></w:r></w:hyperlink>
<w:r><w:t>tail</w:t></w:r>
</w:p>
</w:body></w:document>"""
    write_minimal_docx(
        OUT / "13_formatted_whitespace_only_runs.docx",
        document_xml,
        extra_rels=[
            (
                "rIdL1",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                "https://example.com",
            ),
        ],
    )


# ============================================================
# helper to write a minimal valid docx with custom XML
# ============================================================
def write_minimal_docx(path, document_xml, extra_rels=None, footnotes_xml=None):
    extra_rels = extra_rels or []

    content_types = """<?xml version="1.0"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>"""
    if footnotes_xml:
        content_types += '\n<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
    content_types += "\n</Types>"

    root_rels = """<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    doc_rels_parts = ['<Relationship Id="rIdSt" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>']
    if footnotes_xml:
        doc_rels_parts.append('<Relationship Id="rIdFn" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>')
    for rid, rtype, target in extra_rels:
        if rtype.endswith("/hyperlink"):
            doc_rels_parts.append(f'<Relationship Id="{rid}" Type="{rtype}" Target="{target}" TargetMode="External"/>')
        else:
            doc_rels_parts.append(f'<Relationship Id="{rid}" Type="{rtype}" Target="{target}"/>')
    doc_rels = '<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">' + "".join(doc_rels_parts) + "</Relationships>"

    styles_xml = """<?xml version="1.0"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/></w:style>
<w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/></w:style>
<w:style w:type="paragraph" w:styleId="Heading3"><w:name w:val="heading 3"/></w:style>
<w:style w:type="character" w:styleId="Hyperlink"><w:name w:val="Hyperlink"/></w:style>
</w:styles>"""

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", root_rels)
        z.writestr("word/document.xml", document_xml)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/styles.xml", styles_xml)
        if footnotes_xml:
            z.writestr("word/footnotes.xml", footnotes_xml)


if __name__ == "__main__":
    for name, fn in list(globals().items()):
        if name.startswith("build_") and callable(fn):
            print(f"Building {name}...")
            fn()
    print("\nDONE. Files in", OUT)
